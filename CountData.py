#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import sys

import pandas

reload(sys)
sys.setdefaultencoding("utf-8")

import docx
import win32com.client

WORDSDIRS_DOC = 'WordsFiles_doc'
WORDSDIRS_DOCX = 'WordsFiles_docx'
ABS_WORDSDIRS_DOC = os.path.abspath(os.path.dirname(WORDSDIRS_DOC)) + '\\WordsFiles_doc'
ABS_WORDSDIRS_DOCX = os.path.abspath(os.path.dirname(WORDSDIRS_DOCX)) + '\\WordsFiles_docx'
A_S = list('ABCDEFGHIJKLMNOPQR')
A_V = list('ABCDEFGHIJKLMNOPQRSTUV')


def read_from_word():
    files = os.listdir(WORDSDIRS_DOCX)
    sum_second_table_data_pd = None
    sum_first_table_data_pd = None
    for words_file in files:
        if not os.path.isdir(words_file):
            words_file = ABS_WORDSDIRS_DOCX + '\\' + u"{}".format(words_file.decode('gb2312').encode('utf-8'))
            word_file = docx.Document(words_file)
            first_table = word_file.tables[0]
            index_list = []
            row_data_list = []
            for row in first_table.rows:
                data_cells = row.cells
                index_list.append(data_cells[0].text)
                if not data_cells[0].text[0].isdigit():
                    continue
                row_data = [(data_cell.text.encode('utf-8')) for data_cell in data_cells if '岁' not in data_cell.text]
                map(eval, row_data)
                row_data_list.append(row_data)
            first_table_data_pd = pandas.DataFrame(row_data_list, index=index_list[2:], columns=A_S, dtype=float)
            sum_first_table_data_pd = sum_first_table_data_pd.add(
                first_table_data_pd) if sum_first_table_data_pd is not None \
                else first_table_data_pd
            # second table
            second_table = word_file.tables[1]
            second_index_list = []
            second_row_data_list = []
            for row in second_table.rows:
                second_data_cells = row.cells
                second_index_list.append(second_data_cells[0].text)
                if not second_data_cells[0].text[0].isdigit():
                    continue
                second_row_data = [(data_cell.text.encode('utf-8')) for data_cell in second_data_cells if
                                   '岁' not in data_cell.text]
                map(eval, second_row_data)
                second_row_data_list.append(second_row_data)
            second_table_data_pd = pandas.DataFrame(second_row_data_list, index=second_index_list[3:], columns=A_V,
                                                    dtype=float)
            sum_second_table_data_pd = sum_second_table_data_pd.add(
                second_table_data_pd) if sum_second_table_data_pd is not None \
                else second_table_data_pd
            # pandas.set_option('display.max_columns', None)
            # pandas.set_option('display.max_rows', None)
    writer = pandas.ExcelWriter('tables.xlsx')
    sum_first_table_data_pd.to_excel(writer, sheet_name='Sheet')
    sum_second_table_data_pd.to_excel(writer, sheet_name='Sheet', startrow=11)
    print sum_first_table_data_pd
    print sum_second_table_data_pd
    writer.save()


def doc_to_docx(path):
    if os.path.splitext(path)[1] == ".doc":
        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(path)
        old_path = os.path.splitext(path)[0]
        save_dir = old_path.replace(WORDSDIRS_DOC, WORDSDIRS_DOCX)
        doc.SaveAs(save_dir + ".docx", 12)
        try:
            doc.Close()
            word.Quit()
        except Exception as e:
            print(e)
        path = path + 'x'
    return path


def all_doc_to_docx():
    files = os.listdir(WORDSDIRS_DOC)
    for words_file in files:
        if not os.path.isdir(words_file):
            words_file = ABS_WORDSDIRS_DOC + '\\' + u"{}".format(words_file.decode('gb2312').encode('utf-8'))
            doc_to_docx(words_file)


if __name__ == '__main__':
    # all_doc_to_docx()
    read_from_word()
